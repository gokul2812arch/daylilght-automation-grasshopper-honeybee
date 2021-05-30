try:
    from Tkinter import *
    from Tkinter import filedialog
    import tkFileDialog

except ImportError:
    from tkinter import *
    from tkinter import filedialog

import sys
import math
from docx.shared import Inches
import csv 
import xlrd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from xlrd import open_workbook,XL_CELL_TEXT,XLRDError
from xlrd import open_workbook
import os
import shutil
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from xlrd import open_workbook,XL_CELL_TEXT,XLRDError
from PIL import Image, ImageChops
import numpy as np 
import operator
from os import listdir
from os.path import isfile, join
import glob
import decimal


#global paths 
global folderPath
global templatePath
global wordPath
# global files list 
global png_list
global jpg_list
global xl_list
global csv_list
global e1 
global e2


#get png file names
def import_dir_png(folder):
	png_list = []
	for file in os.listdir(folder):
		if file.endswith('.png'):
			png_list.append(os.path.join(folder, file))
	return (png_list)

#get png file names
def import_dir_jpg(folder):
	jpg_list = []
	for file in os.listdir(folder):
		if file.endswith('.jpg'):
			jpg_list.append(os.path.join(folder, file))
	return (jpg_list)


#get xl file names 
def import_dir_xlsx(folder):
	xl_list = []
	for file in os.listdir(folder):
		if file.endswith('.xlsx'):
			xl_list.append(os.path.join(folder, file))
	return (xl_list)

#get csv file names 
def import_dir_csv(folder):
	csv_list = []
	for file in os.listdir(folder):
		if file.endswith('.csv'):
			csv_list.append(os.path.join(folder, file))
	return (csv_list)

#convert xlsx to csv with same file name 
def csv_from_excel(excel_filename):
    wb = xlrd.open_workbook(excel_filename)
    sh = wb.sheet_by_name('Sheet1')
    csv_filename = excel_filename[:-5] + '.csv'
    your_csv_file = open(csv_filename, 'w')
    wr = csv.writer(your_csv_file)

    for rownum in range(sh.nrows):
        wr.writerow(sh.row_values(rownum))

    your_csv_file.close()

#read csv file to list 
def read_csv(csv_file):
	data = open(csv_file).read()
	dataList = [line.split(',') for line in data.split('\n') if line]
	return dataList

#resize images 
def resize(image):
	width, height = image.size
	final_height = int((5000*height)/width)
	im2 = image.resize((5000,final_height))
	return(im2)

#resize images 
def resize_2(image):
	width, height = image.size
	scale_factor = width/1200
	im3 = image.resize((int(width/scale_factor),int(height/scale_factor)))
	return(im3)

#trim images 
def trim (image):
	bg = Image.new(image.mode, image.size, image.getpixel((0,0)))
	diff = ImageChops.difference(image,bg)
	diff = ImageChops.add(diff, diff, 2.0, -100)
	bbox = diff.getbbox()
	if bbox:
		return image.crop(bbox)

# caption figures 
def Figure(paragraph):
	run = run = paragraph.add_run()
	r = run._r
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')
	r.append(fldChar)
	instrText = OxmlElement('w:instrText')
	instrText.text = ' SEQ Figure * ARABIC'
	r.append(instrText)
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'end')
	r.append(fldChar)

#get caption 
def get_caption(filename):
	a,b = filename.split("_")
	new_a = "Level " + a[6:-4]
	new_b = " " +b[:-4]
	return new_a+new_b

## browse for folder with picture 
def browse_folder():
	global png_list
	global jpg_list
	global xl_list
	global csv_list
	global folderPath

	folderPath = filedialog.askdirectory()
	print(folderPath,"   got folder path")

	#get png files
	try:
		png_list = import_dir_png(folderPath + "/")
		print("got png files")
		print(png_list)
	except:
		print("no png files in the folder")

	#get jpg files 
	try:
		jpg_list = import_dir_jpg(folderPath + "/")
		print("got jpg files")
		print(jpg_list)
	except:
		print("no jpg files in the folder")

	#get xl files 
	try:
		xl_list = import_dir_xlsx(folderPath + "/")
		# runs the csv_from_excel function:
		csv_from_excel(xl_list[0])
		csv_list = import_dir_csv(folderPath + "/")
		print("got xl files")
		print(xl_list,csv_list)
	except:
		print("no xl file present")
	

## browse for folder with picture 
def browse_template():
	global document
	templatePath = filedialog.askopenfilename(filetypes = (("template file", "*.docx"), ("All files", "*")))
	print (templatePath, "get word template")
	# open template 
	document = Document(templatePath)
	print("got template file")

## browse word file
def browse_word():
	global document
	wordPath = filedialog.askopenfilename(filetypes = (("word file", "*.docx"), ("All files", "*")))
	print (wordPath, "get word file")
	# open template 
	document = Document(wordPath)
	print("got word file")

#add table 
def add_table():
	digit_to_round_to = 2
	#read csv file 
	records = read_csv(csv_list[0]) 
	table_height = len(records)
	tabel_width  = len(records[0])
	heading_table = records[0][0]
	table = document.add_table(rows=0, cols=5,)
	table.style = 'TableGrid'
	for c1, c2, c3, c4, c5 in records:
		row_cells = table.add_row().cells

		# round off if number 
		try:
			n1 = str(round(float(c1),digit_to_round_to))
		except:
			n1 = c1

		try:
			n2 = str(round(float(c2),digit_to_round_to))
		except:
			n2 = c2

		try:
			n3 = str(round(float(c3),digit_to_round_to))
		except:
			n3 = c3

		try:
			n4 = str(round(float(c4),digit_to_round_to))
		except:
			n4 = c4

		try:
			n5 = str(round(float(c5),digit_to_round_to))
		except:
			n5 = c5 

			
		row_cells[0].text = n1
		row_cells[1].text = n2
		row_cells[2].text = n3
		row_cells[3].text = n4
		row_cells[4].text = n5
	row = table.rows[0]
	merge1 = row.cells[0]
	merge2 = row.cells[4]
	merge1.merge(merge2)
	row.cells[0].text = heading_table
	print("table added")

#insert images 
def insert_Images():
	#trim all images
	global folderPath
	caption_remove = len(folderPath)
	for i in png_list:
		
		#to backup jpg 
		jp = Image.open(i)
		rgb_jp = jp.convert('RGB')
		rgb_jp.save(i[:-4]+".jpg",dpi=(300,300))

		# IMAG TRIM 
		im = Image.open(i)
		im = resize_2(trim(resize(im)))
		im.save(i,dpi=(300,300))
	for i in png_list:
		#add picture
		document.add_picture(i) # ,width=Inches(5)
		# center align
		last_paragraph = document.paragraphs[-1] 
		last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

		#caption 
		c1 = document.add_paragraph(get_caption(i[caption_remove:]) ,style='Caption')
		Figure (c1)
		# center align
		last_paragraph = document.paragraphs[-1] 
		last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	print("images inserted")

def save_document():
	document.save('report1.docx')
	print("document saved")

def jpg_png():
	global folderPath
	global png_list
	global jpg_list
	global xl_list
	global csv_list
	global folderPath

		#get png files
	try:
		png_list = import_dir_png(folderPath + "/")
		print("got png files")
		print(png_list)
	except:
		print("no png files in the folder")

	#get jpg files 
	try:
		jpg_list = import_dir_jpg(folderPath + "/")
		print("got jpg files")
		print(jpg_list)
	except:
		print("no jpg files in the folder")

	for w in jpg_list:
		im2 = Image.open(w)
		im2.save(w[:-4]+".png",dpi=(300,300))
	print ("backup restored ")

def add_heading():
	global e1
	global e2
	heading = e1.get()
	sub_heading = e2.get()
	document.add_heading(heading, level=2)
	document.add_heading(sub_heading, level=3)

	#space
	enter = document.add_paragraph(' ' ,style='Normal')
	print("heading added")



#main
master = Tk()
line_row = 0

# browse for folder with picture 
Button(master, text='browse', command=browse_folder).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="browse folder with files").grid(row=line_row)
line_row = line_row + 1

# break
Label(master, text="-----------------------------").grid(row=line_row)
Label(master, text="----------------------------------------------------------").grid(row=line_row, column=1)
line_row = line_row + 1

# browse template and create word file
Button(master, text='browse', command=browse_template).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="browse word template").grid(row=line_row)
line_row = line_row + 1

# browse existing word file 
Button(master, text='browse', command=browse_word).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="browse existing word file").grid(row=line_row)
line_row = line_row + 1

# break
Label(master, text="-----------------------------").grid(row=line_row)
Label(master, text="----------------------------------------------------------").grid(row=line_row, column=1)
line_row = line_row + 1

#text boxes for threshold
Label(master, text="add heading").grid(row=line_row)
line_row = line_row + 1

Label(master, text="heading").grid(row=line_row)
e1 = Entry(master)
e1.grid(row=line_row, column=1)
e1.insert(100,"1.2  Option 1 Analysis & Calculations")
line_row = line_row + 1

Label(master, text="sub heading").grid(row=line_row)
e2 = Entry(master)
e2.insert(100,"Spatial Daylight Autonomy Results")
e2.grid(row=line_row, column=1)
line_row = line_row + 1


# browse existing word file 
Button(master, text='add', command=add_heading).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="click to add heading").grid(row=line_row)
line_row = line_row + 1

# break
Label(master, text="-----------------------------").grid(row=line_row)
Label(master, text="----------------------------------------------------------").grid(row=line_row, column=1)
line_row = line_row + 1

# insert table
Button(master, text='add', command=add_table).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="add table from excel file ").grid(row=line_row)
line_row = line_row + 1

#insert images 
Button(master, text='add', command=insert_Images).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="trim and insert images save backup as jpg").grid(row=line_row)
line_row = line_row + 1

# break 
Label(master, text="-----------------------------").grid(row=line_row)
Label(master, text="----------------------------------------------------------").grid(row=line_row, column=1)
line_row = line_row + 1

#insert images 
Button(master, text='restore', command=jpg_png).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="convert jpg to png").grid(row=line_row)
line_row = line_row + 1

# break
Label(master, text="-----------------------------").grid(row=line_row)
Label(master, text="----------------------------------------------------------").grid(row=line_row, column=1)
line_row = line_row + 1

#save document 
Button(master, text='save', command=save_document).grid(row=line_row, column=1, sticky=W, pady=4)
Label(master, text="save document").grid(row=line_row)
line_row = line_row + 1

#exit 
mainloop()